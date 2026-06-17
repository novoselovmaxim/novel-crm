"""Generate CP (commercial proposal) .docx for a company."""

import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MEDIA_DIR = os.path.join(os.path.dirname(__file__), "media")

O = RGBColor(0xFE, 0x5B, 0x24)
D = RGBColor(0x1A, 0x1A, 0x2E)
G = RGBColor(0x66, 0x66, 0x66)
L = RGBColor(0x99, 0x99, 0x99)
W = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is None:
            val = {'sz': '0', 'color': 'FFFFFF', 'val': 'none', 'space': '1'}
        tag = tcBorders.find(qn(f'w:{edge}'))
        if tag is None:
            tag = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
            tcBorders.append(tag)
        tag.set(qn('w:val'), val.get('val', 'none'))
        tag.set(qn('w:sz'), str(val.get('sz', '0')))
        tag.set(qn('w:color'), val.get('color', 'FFFFFF'))
        tag.set(qn('w:space'), str(val.get('space', '1')))


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}></w:tcMar>')
        tcPr.append(tcMar)
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        tag = tcMar.find(qn(f'w:{edge}'))
        if tag is None:
            tag = parse_xml(f'<w:{edge} {nsdecls("w")}/>')
            tcMar.append(tag)
        tag.set(qn('w:w'), str(val))
        tag.set(qn('w:type'), 'dxa')


def _set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_vertical_align(cell, align="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
        tcPr.append(vAlign)
    else:
        vAlign.set(qn('w:val'), align)


def _add_run(paragraph, text, size=12, color=D, bold=False, font="Arial"):
    run = paragraph.add_run(text)
    run.font.size = Pt(size / 2)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}" w:eastAsia="{font}"/>')
        rPr.append(rFonts)
    return run


def _add_paragraph(doc_or_cell, text_runs=None, alignment=None, spacing_after=60, spacing_before=0, line=248, indent_left=None):
    p = doc_or_cell.add_paragraph()
    
    pPr = p._p.get_or_add_pPr()
    pPr_spacing = pPr.find(qn('w:spacing'))
    if pPr_spacing is None:
        pPr_spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(pPr_spacing)
    pPr_spacing.set(qn('w:after'), str(spacing_after))
    pPr_spacing.set(qn('w:before'), str(spacing_before))
    pPr_spacing.set(qn('w:line'), str(line))
    
    if alignment:
        pPr.alignment = alignment
    
    if indent_left:
        pPr_ind = pPr.find(qn('w:ind'))
        if pPr_ind is None:
            pPr_ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
            pPr.append(pPr_ind)
        pPr_ind.set(qn('w:left'), str(indent_left))
    
    if text_runs:
        for tr_data in text_runs:
            _add_run(p, **tr_data)
    
    return p


def _add_image_to_cell(cell, image_path, width_emu, height_emu):
    p = _add_paragraph(cell, alignment=WD_ALIGN_PARAGRAPH.CENTER, line=240)
    run = p.add_run()
    run.add_picture(image_path, width=Emu(width_emu), height=Emu(height_emu))


def _add_image_to_cell_left(cell, image_path, width_emu, height_emu):
    p = _add_paragraph(cell, alignment=WD_ALIGN_PARAGRAPH.LEFT, line=240)
    run = p.add_run()
    run.add_picture(image_path, width=Emu(width_emu), height=Emu(height_emu))


def _make_cell(cell, borders_all=None, borders=None, margins=None, shading=None, vertical_align="center"):
    if borders_all:
        _set_cell_border(cell, top=borders_all, bottom=borders_all, left=borders_all, right=borders_all)
    if borders:
        _set_cell_border(cell, **borders)
    if margins:
        _set_cell_margins(cell, **margins)
    if shading:
        _set_cell_shading(cell, shading)
    _set_cell_vertical_align(cell, vertical_align)


def _add_hyperlink(paragraph, text, url, size=20, color=G, bold=False, font="Arial"):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink_xml = (
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}>'
        f'<w:r><w:rPr>'
        f'<w:color w:val="{color}"/>'
        f'<w:sz w:val="{size}"/>'
        f'<w:u w:val="single"/>'
        f'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}" w:eastAsia="{font}"/>'
        f'</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:hyperlink>'
    )
    paragraph._p.append(parse_xml(hyperlink_xml))


def _empty_paragraph(doc_or_cell, after=60):
    return _add_paragraph(doc_or_cell, spacing_after=after, spacing_before=0, line=240)


NB = {'val': 'none', 'sz': '0', 'color': 'FFFFFF', 'space': '1'}


def _set_paragraph_spacing(p, after=0, before=0, line=240):
    pPr = p._p.get_or_add_pPr()
    pPr_spacing = pPr.find(qn('w:spacing'))
    if pPr_spacing is None:
        pPr_spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(pPr_spacing)
    pPr_spacing.set(qn('w:after'), str(after))
    pPr_spacing.set(qn('w:before'), str(before))
    pPr_spacing.set(qn('w:line'), str(line))
    return pPr


def _add_bottom_border(pPr, sz="2", color="DDDDDD"):
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:color="{color}" w:space="1"/></w:pBdr>')
    pPr.append(pBdr)


def _add_orange_line(pPr):
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:color="FE5B24" w:space="1"/></w:pBdr>')
    pPr.append(pBdr)


def _clear_cell(cell):
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def generate_cp(company_name="", lpr_name="", lpr_phone="", lpr_firstname=""):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Emu(11906 * 914400 // 1440)
    section.page_height = Emu(16838 * 914400 // 1440)
    section.top_margin = Cm(0.88)
    section.bottom_margin = Cm(0.88)
    section.left_margin = Cm(1.06)
    section.right_margin = Cm(1.06)

    # === TOP ORANGE LINE ===
    line_p = doc.add_paragraph()
    pPr = _set_paragraph_spacing(line_p, after=0, before=0, line=240)
    _add_orange_line(pPr)

    # === HEADER ===
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Cm(7.06)
    header_table.columns[1].width = Cm(11.83)

    # Left cell: logo + text
    left_cell = header_table.cell(0, 0)
    _make_cell(left_cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB},
               margins={'left': 120, 'right': 60, 'top': 10, 'bottom': 10})

    nested = left_cell.add_table(rows=1, cols=2)
    nested.columns[0].width = Cm(1.41)
    nested.columns[1].width = Cm(5.0)

    img_cell = nested.cell(0, 0)
    _clear_cell(img_cell)
    _make_cell(img_cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB},
               margins={'left': 0, 'right': 0, 'top': 0, 'bottom': 0})
    logo_int_path = os.path.join(MEDIA_DIR, "intpay_logo.jpg")
    if os.path.exists(logo_int_path):
        _add_image_to_cell_left(img_cell, logo_int_path, 685800, 685800)

    txt_cell = nested.cell(0, 1)
    _clear_cell(txt_cell)
    _make_cell(txt_cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB},
               margins={'left': 0, 'right': 0, 'top': 0, 'bottom': 0})
    _add_paragraph(txt_cell, [{'text': 'ИНТПЭЙ', 'size': 32, 'color': O, 'bold': True}],
                   alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Right cell: novel + amf logos
    right_cell = header_table.cell(0, 1)
    _make_cell(right_cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB},
               margins={'left': 0, 'right': 60, 'top': 20, 'bottom': 10})

    logo_novel_path = os.path.join(MEDIA_DIR, "logo_novel.png")
    logo_amf_path = os.path.join(MEDIA_DIR, "logo-1.png")

    right_p = _add_paragraph(right_cell, alignment=WD_ALIGN_PARAGRAPH.RIGHT, spacing_after=0, line=240, spacing_before=0)
    if os.path.exists(logo_novel_path):
        right_run = right_p.add_run()
        right_run.add_picture(logo_novel_path, width=Emu(762000), height=Emu(228600))
    right_space = right_p.add_run('  ')
    right_space.font.size = Pt(4)
    if os.path.exists(logo_amf_path):
        right_run2 = right_p.add_run()
        right_run2.add_picture(logo_amf_path, width=Emu(666750), height=Emu(342900))

    # === EMPTY PARAGRAPH BETWEEN HEADER AND CONTENT (after=120) ===
    spacer_p = doc.add_paragraph()
    _set_paragraph_spacing(spacer_p, after=120, before=0, line=240)

    # === CONTENT SECTION ===
    content_table = doc.add_table(rows=1, cols=1)
    content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    content_table.columns[0].width = Cm(18.88)

    content_cell = content_table.cell(0, 0)
    _make_cell(content_cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB},
               margins={'left': 1134, 'right': 1134, 'top': 20, 'bottom': 20})

    # Remove the default empty paragraph that comes with a new cell
    for p in content_cell.paragraphs:
        p._p.getparent().remove(p._p)

    # P0: Company info
    ci = _add_paragraph(content_cell, alignment=WD_ALIGN_PARAGRAPH.RIGHT, spacing_after=40, line=248)
    _add_run(ci, 'Компания: ', 22, L)
    _add_run(ci, company_name, 22, D, True)
    ci.add_run().add_break()
    _add_run(ci, lpr_name, 22, D, True)
    ci.add_run().add_break()
    _add_run(ci, 'Тел: ', 22, L)
    _add_run(ci, lpr_phone, 22, D, True)

    # P1: Empty (after=4)
    _empty_paragraph(content_cell, after=4)

    # P2: Salutation
    _add_paragraph(content_cell, [
        {'text': 'Уважаемый ', 'size': 24, 'color': D, 'bold': False},
        {'text': f'{lpr_firstname}!', 'size': 24, 'color': O, 'bold': True},
    ], spacing_before=40, spacing_after=40, line=248)

    # P3: Empty (after=120)
    _empty_paragraph(content_cell, after=120)

    # P4: Intro
    _add_paragraph(content_cell, [
        {'text': 'Мы — ', 'size': 24, 'color': D, 'bold': False},
        {'text': 'ИНТПЭЙ', 'size': 24, 'color': O, 'bold': True},
        {'text': ', платёжное подразделение международного холдинга ', 'size': 24, 'color': D, 'bold': False},
        {'text': 'NOVEL GROUP', 'size': 24, 'color': D, 'bold': True},
        {'text': ': ', 'size': 24, 'color': D, 'bold': False},
        {'text': '24 года', 'size': 24, 'color': D, 'bold': True},
        {'text': ' на рынке, портфель недвижимости >100 000 м², телеканалы «Моя Планета» и «Наука» с аудиторией >50 млн человек, собственная платёжная система с выпуском карт VISA.', 'size': 24, 'color': D, 'bold': False},
    ], spacing_after=40, line=248)

    # P5: Empty (after=2)
    _empty_paragraph(content_cell, after=2)

    # P6: AMF paragraph
    _add_paragraph(content_cell, [
        {'text': 'Партнёрство с ', 'size': 24, 'color': D, 'bold': False},
        {'text': 'Арабским валютным фондом (AMF)', 'size': 24, 'color': D, 'bold': True},
        {'text': ' гарантирует полную юридическую чистоту каждого перевода.', 'size': 24, 'color': D, 'bold': False},
    ], spacing_after=40, line=248)

    # P7: Empty (after=4)
    _empty_paragraph(content_cell, after=4)

    # P8: НАШИ ПРЕИМУЩЕСТВА
    section_p = _add_paragraph(content_cell, [{'text': 'НАШИ ПРЕИМУЩЕСТВА', 'size': 22, 'color': O, 'bold': True}],
                                alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0, spacing_before=0, line=240)
    _add_bottom_border(section_p._p.get_or_add_pPr())

    # P9: Empty (after=2)
    _empty_paragraph(content_cell, after=2)

    # P10: Empty (after=4)
    _empty_paragraph(content_cell, after=4)

    # Advantages 2x2 table
    adv_items = [
        ('Экономия до 70%', 'Комиссия от 0,5% вместо банковских 2–5%.'),
        ('Скорость 1–3 дня', 'Срочно за 24 часа.'),
        ('Валютный контроль', 'Полный пакет документов для ЦБ и вашего банка.'),
        ('Любые направления', 'Европа, Азия, ОАЭ, Китай, Великобритания, США.'),
    ]

    adv_table = content_cell.add_table(rows=2, cols=3)
    adv_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    be = {'val': 'single', 'sz': '1', 'color': 'EEEEEE', 'space': '1'}
    bl = {'val': 'single', 'sz': '4', 'color': 'FE5B24', 'space': '1'}

    for i, (title, desc) in enumerate(adv_items):
        row_idx = i // 2
        col_idx = (i % 2) * 2
        cell = adv_table.cell(row_idx, col_idx)
        _clear_cell(cell)
        _make_cell(cell, borders={'top': be, 'bottom': be, 'left': bl, 'right': be},
                   margins={'top': 20, 'bottom': 20, 'left': 80, 'right': 80})
        _add_paragraph(cell, [{'text': title, 'size': 22, 'color': O, 'bold': True}], spacing_after=40, line=248)
        _add_paragraph(cell, [{'text': desc, 'size': 20, 'color': G, 'bold': False}], spacing_after=40, line=248)

        if col_idx == 0:
            sep_cell = adv_table.cell(row_idx, 1)
            _make_cell(sep_cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB})

    # P11: СХЕМА РАБОТЫ
    section_p2 = _add_paragraph(content_cell, [{'text': 'СХЕМА РАБОТЫ', 'size': 22, 'color': O, 'bold': True}],
                                alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0, spacing_before=0, line=240)
    _add_bottom_border(section_p2._p.get_or_add_pPr())

    # P12: Empty (after=2)
    _empty_paragraph(content_cell, after=2)

    # P13: Empty (after=4)
    _empty_paragraph(content_cell, after=4)

    # Steps table — three separate paragraphs per cell
    steps_data = [
        ('01', 'Заявка', '→ ответ за 30 минут'),
        ('02', 'Договор', '→ тариф под ваш объём'),
        ('03', 'Перевод', '→ зачисление за 1–3 дня'),
    ]

    steps_table = content_cell.add_table(rows=1, cols=5)
    steps_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (num, title, desc) in enumerate(steps_data):
        cell = steps_table.cell(0, i * 2)
        _clear_cell(cell)
        _make_cell(cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB})

        _add_paragraph(cell, [{'text': num, 'size': 20, 'color': O, 'bold': True}],
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=40, line=248)
        _add_paragraph(cell, [{'text': title, 'size': 22, 'color': D, 'bold': True}],
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=40, line=248)
        _add_paragraph(cell, [{'text': desc, 'size': 18, 'color': G, 'bold': False}],
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=40, line=248)

        if i < 2:
            sep = steps_table.cell(0, i * 2 + 1)
            _make_cell(sep, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB})

    # P14: ПОЧЕМУ НАМ ДОВЕРЯЮТ
    section_p3 = _add_paragraph(content_cell, [{'text': 'ПОЧЕМУ НАМ ДОВЕРЯЮТ', 'size': 22, 'color': O, 'bold': True}],
                                alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0, spacing_before=0, line=240)
    _add_bottom_border(section_p3._p.get_or_add_pPr())

    # P15: Empty (after=2)
    _empty_paragraph(content_cell, after=2)

    # P16-P18: Trust items
    trust_items = [
        '24 года на рынке',
        'Партнёрство с AMF',
        'Работаем через крупнейшие банки-партнёры (см. логотипы)',
    ]
    for item in trust_items:
        _add_paragraph(content_cell, [
            {'text': '✓ ', 'size': 22, 'color': O, 'bold': True},
            {'text': item, 'size': 22, 'color': D, 'bold': False},
        ], spacing_after=20, line=248)

    # P19: Empty (after=4)
    _empty_paragraph(content_cell, after=4)

    # P20: БАНКИ-ПАРТНЁРЫ
    section_p4 = _add_paragraph(content_cell, [{'text': 'БАНКИ-ПАРТНЁРЫ', 'size': 22, 'color': O, 'bold': True}],
                                alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0, spacing_before=0, line=240)
    _add_bottom_border(section_p4._p.get_or_add_pPr())

    # P21: Empty (after=2)
    _empty_paragraph(content_cell, after=2)

    # P22: Empty (after=4)
    _empty_paragraph(content_cell, after=4)

    # Banks logos
    banks_table = content_cell.add_table(rows=1, cols=3)
    banks_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    bank_logo_specs = [
        (os.path.join(MEDIA_DIR, "alfa_bank.png"), 952500, 152400),
        (os.path.join(MEDIA_DIR, "sovcombank.png"), 1143000, 152400),
        (os.path.join(MEDIA_DIR, "mts_bank.png"), 952500, 152400),
    ]
    for ci, (bpath, w_emu, h_emu) in enumerate(bank_logo_specs):
        cell = banks_table.cell(0, ci)
        _clear_cell(cell)
        _make_cell(cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB})
        if os.path.exists(bpath):
            _add_image_to_cell(cell, bpath, w_emu, h_emu)

    # P23: P.S.
    _add_paragraph(content_cell, [
        {'text': 'P.S. ', 'size': 22, 'color': O, 'bold': True},
        {'text': 'Пришлите один платёж — сделаем перевод за 0,5% и покажем разницу. Отвечаем на заявку в течение 30 минут.', 'size': 22, 'color': D, 'bold': False},
    ], spacing_after=0, spacing_before=0, line=248)

    # === BOTTOM ORANGE LINE ===
    line_p2 = doc.add_paragraph()
    pPr2 = _set_paragraph_spacing(line_p2, after=0, before=0, line=240)
    _add_orange_line(pPr2)

    # === FOOTER ===
    footer_table = doc.add_table(rows=1, cols=1)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.columns[0].width = Cm(18.88)

    footer_cell = footer_table.cell(0, 0)
    _clear_cell(footer_cell)
    _make_cell(footer_cell, borders={'top': NB, 'bottom': NB, 'left': NB, 'right': NB},
               margins={'left': 200, 'right': 200, 'top': 25, 'bottom': 25})

    _add_paragraph(footer_cell, [{'text': 'Свяжитесь с нами для индивидуального тарифа:', 'size': 22, 'color': D, 'bold': True}],
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=40, line=248)

    contact_p = _add_paragraph(footer_cell, alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=40, line=248)
    _add_hyperlink(contact_p, 'Сайт — intpaypro.ru', 'https://intpaypro.ru', 20, G)
    _add_run(contact_p, '  ·  ', 20, L)
    _add_hyperlink(contact_p, 'E-mail — info@intpaypro.ru', 'mailto:info@intpaypro.ru', 20, G)
    _add_run(contact_p, '  ·  ', 20, L)
    tg_path = os.path.join(MEDIA_DIR, "tg_logo.png")
    if os.path.exists(tg_path):
        tg_run = contact_p.add_run('  ')
        tg_run.font.size = Pt(4)
        tg_run2 = contact_p.add_run()
        tg_run2.add_picture(tg_path, width=Emu(133350), height=Emu(171450))
        _add_hyperlink(contact_p, ' @in_veritate', 'https://t.me/in_veritate', 20, G)

    # Clean up auto-generated empty paragraphs in content cell
    for p in list(content_cell.paragraphs):
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None:
            sp = pPr.find(qn('w:spacing'))
            if sp is not None:
                continue
        if not p.text.strip():
            p._p.getparent().remove(p._p)

    # Write to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _img_to_data_uri(path):
    """Read image and return data URI for embedding in HTML email."""
    if not os.path.exists(path):
        return None
    import base64
    with open(path, "rb") as f:
        data = f.read()
    ext = os.path.splitext(path)[1].lower()
    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg"}.get(
        ext.lstrip("."), "image/png"
    )
    return f"data:{mime};base64,{base64.b64encode(data).decode()}"


def generate_cp_html(company_name="", lpr_name="", lpr_phone="", lpr_firstname=""):
    """Generate CP as an HTML string suitable for email body."""
    intpay_logo = _img_to_data_uri(os.path.join(MEDIA_DIR, "intpay_logo.jpg"))
    novel_logo = _img_to_data_uri(os.path.join(MEDIA_DIR, "logo_novel.png"))
    amf_logo = _img_to_data_uri(os.path.join(MEDIA_DIR, "logo-1.png"))
    alfa_logo = _img_to_data_uri(os.path.join(MEDIA_DIR, "alfa_bank.png"))
    sovcombank_logo = _img_to_data_uri(os.path.join(MEDIA_DIR, "sovcombank.png"))
    mts_logo = _img_to_data_uri(os.path.join(MEDIA_DIR, "mts_bank.png"))
    tg_logo = _img_to_data_uri(os.path.join(MEDIA_DIR, "tg_logo.png"))

    O = "#FE5B24"
    D = "#1A1A2E"
    G = "#666666"
    L = "#999999"

    adv_rows = ""
    for title, desc in [
        ("Экономия до 70%", "Комиссия от 0,5% вместо банковских 2–5%."),
        ("Скорость 1–3 дня", "Срочно за 24 часа."),
        ("Валютный контроль", "Полный пакет документов для ЦБ и вашего банка."),
        ("Любые направления", "Европа, Азия, ОАЭ, Китай, Великобритания, США."),
    ]:
        adv_rows += f"""<tr>
          <td style="padding:10px 20px 10px 15px;border-bottom:1px solid #eee;border-left:3px solid {O};vertical-align:top;width:50%">
            <div style="font-size:14px;color:{O};font-weight:bold;margin-bottom:4px">{title}</div>
            <div style="font-size:12px;color:{G}">{desc}</div>
          </td>
        </tr>"""

    steps_html = ""
    for num, title, desc in [
        ("01", "Заявка", "→ ответ за 30 минут"),
        ("02", "Договор", "→ тариф под ваш объём"),
        ("03", "Перевод", "→ зачисление за 1–3 дня"),
    ]:
        steps_html += f"""<td style="padding:10px 15px;vertical-align:top;text-align:center;width:33%">
          <div style="font-size:12px;color:{O};font-weight:bold">{num}</div>
          <div style="font-size:14px;color:{D};font-weight:bold;margin:4px 0">{title}</div>
          <div style="font-size:11px;color:{G}">{desc}</div>
        </td>"""

    trust_list = "".join(
        f'<tr><td style="padding:4px 0;font-size:13px;color:{D}">'
        f'<span style="color:{O};font-weight:bold">✓ </span>{item}</td></tr>'
        for item in [
            "24 года на рынке",
            "Партнёрство с AMF",
            "Работаем через крупнейшие банки-партнёры",
        ]
    )

    bank_logos_html = ""
    for logo_data in [alfa_logo, sovcombank_logo, mts_logo]:
        if logo_data:
            bank_logos_html += f"""<td style="padding:10px;text-align:center;width:33%">
              <img src="{logo_data}" style="max-height:40px;max-width:120px" alt="bank"/>
            </td>"""

    logo_block = ""
    if intpay_logo:
        logo_block += f'<img src="{intpay_logo}" style="height:36px;vertical-align:middle" alt="ИНТПЭЙ"/>'

    # Build the full HTML email
    html = f"""<!DOCTYPE html>
<html>
<head><meta charset="utf-8"></head>
<body style="margin:0;padding:0;background:#f4f4f4;font-family:Arial,Helvetica,sans-serif">
<table align="center" cellpadding="0" cellspacing="0" style="width:600px;max-width:100%;background:#ffffff">
  <tr>
    <td style="padding:15px 20px;border-bottom:3px solid {O}">
      <table width="100%" cellpadding="0" cellspacing="0"><tr>
        <td style="vertical-align:middle">{logo_block}</td>
        <td style="text-align:right;vertical-align:middle">
          {f'<img src="{novel_logo}" style="height:16px;vertical-align:middle" alt="NOVEL"/>' if novel_logo else ""}
          {f'<img src="{amf_logo}" style="height:24px;vertical-align:middle;margin-left:8px" alt="AMF"/>' if amf_logo else ""}
        </td>
      </tr></table>
    </td>
  </tr>
  <tr>
    <td style="padding:20px 30px">
      <div style="text-align:right;font-size:13px;color:{L};margin-bottom:4px">Компания: <strong style="color:{D}">{company_name}</strong></div>
      <div style="text-align:right;font-size:13px;color:{D};margin-bottom:16px">
        <strong>{lpr_name}</strong><br>
        <span style="color:{L}">Тел:</span> <strong>{lpr_phone}</strong>
      </div>

      <p style="font-size:14px;color:{D};margin:8px 0">
        Уважаемый <strong style="color:{O}">{lpr_firstname}!</strong>
      </p>

      <p style="font-size:14px;color:{D};line-height:1.5;margin:12px 0">
        Мы — <strong style="color:{O}">ИНТПЭЙ</strong>, платёжное подразделение международного холдинга
        <strong>NOVEL GROUP</strong>: <strong>24 года</strong> на рынке, портфель недвижимости &gt;100 000 м²,
        телеканалы «Моя Планета» и «Наука» с аудиторией &gt;50 млн человек, собственная платёжная система
        с выпуском карт VISA.
      </p>

      <p style="font-size:14px;color:{D};line-height:1.5;margin:12px 0">
        Партнёрство с <strong>Арабским валютным фондом (AMF)</strong> гарантирует полную юридическую
        чистоту каждого перевода.
      </p>

      <div style="text-align:center;font-size:13px;color:{O};font-weight:bold;text-transform:uppercase;padding-bottom:4px;border-bottom:1px solid #ddd;margin-bottom:12px">
        НАШИ ПРЕИМУЩЕСТВА
      </div>
      <table width="100%" cellpadding="0" cellspacing="0">{adv_rows}</table>

      <div style="text-align:center;font-size:13px;color:{O};font-weight:bold;text-transform:uppercase;padding-bottom:4px;border-bottom:1px solid #ddd;margin:16px 0 12px">
        СХЕМА РАБОТЫ
      </div>
      <table width="100%" cellpadding="0" cellspacing="0"><tr>{steps_html}</tr></table>

      <div style="text-align:center;font-size:13px;color:{O};font-weight:bold;text-transform:uppercase;padding-bottom:4px;border-bottom:1px solid #ddd;margin:16px 0 12px">
        ПОЧЕМУ НАМ ДОВЕРЯЮТ
      </div>
      <table cellpadding="0" cellspacing="0">{trust_list}</table>

      <div style="text-align:center;font-size:13px;color:{O};font-weight:bold;text-transform:uppercase;padding-bottom:4px;border-bottom:1px solid #ddd;margin:16px 0 12px">
        БАНКИ-ПАРТНЁРЫ
      </div>
      <table width="100%" cellpadding="0" cellspacing="0"><tr>{bank_logos_html}</tr></table>

      <p style="font-size:13px;color:{D};line-height:1.5;margin:16px 0">
        <strong style="color:{O}">P.S.</strong> Пришлите один платёж — сделаем перевод за 0,5% и покажем
        разницу. Отвечаем на заявку в течение 30 минут.
      </p>
    </td>
  </tr>
  <tr>
    <td style="border-top:3px solid {O};padding:15px 20px;text-align:center">
      <p style="font-size:12px;color:{D};font-weight:bold;margin:0 0 8px">
        Свяжитесь с нами для индивидуального тарифа:
      </p>
      <p style="font-size:11px;color:{G};margin:0">
        Сайт — intpaypro.ru &nbsp;·&nbsp; E-mail — info@intpaypro.ru
        {f'&nbsp;&nbsp;<img src="{tg_logo}" style="height:14px;vertical-align:middle" alt="TG"/> @in_veritate' if tg_logo else ""}
      </p>
    </td>
  </tr>
</table>
</body>
</html>"""

    return html
